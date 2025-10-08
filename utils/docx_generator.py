# utils/docx_generator.py
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === Fayl nomini xavfsiz shaklga keltirish ===
def _sanitize_filename(text: str) -> str:
    safe = re.sub(r"[^\w\- ]+", "", text, flags=re.UNICODE).strip().replace(" ", "_")
    return safe or "konspekt"

# === DOCX yaratish funksiyasi ===
def create_docx(text: str, filename: str = "konspekt.docx", title: str = None):
    safe_filename = _sanitize_filename(os.path.splitext(filename)[0]) + ".docx"
    if len(safe_filename) > 100:
        safe_filename = safe_filename[:100] + ".docx"

    doc = Document()

    # === Umumiy shrift va stil ===
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # === Sarlavha (Title) ===
    if title:
        title_p = doc.add_paragraph(title)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 51, 153)
        doc.add_paragraph("")  # bo‘sh qator

    # === Asosiy matnni formatlash ===
    headers = [
        "Mavzu", "Maqsad", "Kutilayotgan", "Jihoz", "Asosiy",
        "Yangi mavzu", "Qo‘shimcha", "Baholash", "Uyga vazifa",
        "Misollar", "Formulalar", "Darsning borishi", "Amaliy mashqlar"
    ]

    for line in text.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")  # bo‘sh qator
            continue

        # Sarlavhalarni aniqlash
        is_heading = any(line.lower().startswith(h.lower()) for h in headers) \
                     or re.match(r"^\d+[\.\)]\s", line) \
                     or line.endswith(":")

        p = doc.add_paragraph()
        run = p.add_run(line)

        if is_heading:
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.space_after = Pt(6)
            p.space_before = Pt(8)
        else:
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.space_after = Pt(4)

    # === Avtomatik sahifa chetlari va oraliqlarni sozlash ===
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(60)
        section.bottom_margin = Pt(60)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    doc.save(safe_filename)
    return safe_filename


# === Foydalanuvchi nomiga mos fayl yaratish ===
def create_named_docx(text: str, subject: str, topic: str, user_id: int, mode: str = "konspekt"):
    """
    mode = "konspekt" yoki "dars_ishlanma"
    """
    topic_clean = "Kop_mavzular" if "\n" in topic or len(topic) > 50 else _sanitize_filename(topic)
    base_name = f"{user_id}_{_sanitize_filename(subject)}_{topic_clean}_{mode}.docx"

    title = f"{subject} — {topic}" if mode == "konspekt" else f"{subject} ({topic}) — Dars ishlanma"
    return create_docx(text, filename=base_name, title=title)


# === Preview (foydalanuvchiga ko‘rsatish uchun 20%) ===
def get_preview(text: str, percent: int = 20):
    lines = text.splitlines()
    preview_len = max(1, len(lines) * percent // 100)
    return "\n".join(lines[:preview_len])
