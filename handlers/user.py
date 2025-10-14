# handlers/user.py
from aiogram import Router, types, F
from aiogram.filters import Command
import pandas as pd
from docx import Document
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardMarkup, InlineKeyboardButton
from utils.db import (
    add_user, is_premium, is_blocked, save_history,
    set_state, get_state, set_subject, get_subject,
    set_grade, get_grade, add_payment,
    get_free_uses, increment_free_use
)
from utils.openai_api import (
    generate_conspect, generate_lesson_plan, generate_methodical_advice, analyze_teaching_problem
)
from utils.docx_generator import create_named_docx, get_preview
from config import ADMIN_ID
import os, logging

router = Router()
logger = logging.getLogger(__name__)

# === Asosiy menyu ===
def main_menu():
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📄 Yangi Konspekt"), KeyboardButton(text="📘 Dars ishlanma yaratish")],
            [KeyboardButton(text="📙 Metodik maslahat"), KeyboardButton(text="📂 Mening konspektlarim")],
            [KeyboardButton(text="🪄 Muammoni tahlil qilish"), KeyboardButton(text="📤 Excel fayldan konspekt yaratish")]
        ],
        resize_keyboard=True
    )


# === START ===
@router.message(Command("start"))
async def start_handler(msg: types.Message):
    add_user(msg.from_user.id, msg.from_user.username)
    if is_blocked(msg.from_user.id):
        return await msg.answer("⛔ Sizning profilingiz bloklangan.")
    await msg.answer(
        "🎓 Assalomu alaykum!\n\n"
        "Bu bot sizga o‘qituvchilar uchun tayyor KONSPEKT, DARS ISHLANMA va METODIK MASLAHATLAR yaratib beradi. 🤖\n\n"
        "🎁 Har bir yangi foydalanuvchi 3 marta bepul foydalanadi!\n"
        "🔒 Shundan so‘ng, barcha xizmatlardan foydalanish uchun 15 000 UZS to‘lov qilinadi.\n\n"
        "Boshlash uchun quyidagilardan birini tanlang 👇",
        reply_markup=main_menu()
    )


# === Limit tekshiruvi ===
async def check_limit(uid: int, msg: types.Message):
    if uid == ADMIN_ID:
        return True
    if is_premium(uid):
        return True

    free_uses = get_free_uses(uid)
    if free_uses < 3:
        increment_free_use(uid)
        await msg.answer(f"🎁 Bepul foydalanish: {free_uses + 1}/3")
        return True
    else:
        await msg.answer(
            "🎁 Sizning 3 ta bepul imkoniyatingiz tugadi.\n"
            "🔐 Xizmatdan foydalanishni davom ettirish uchun 15 000 UZS to‘lov qiling.\n"
            "💳 Karta: <code>9860 6067 4424 9933</code>\n"
            "📸 To‘lov qilib bo‘lgandan so‘ng, chek rasmini shu botga yuboring — admin tasdiqlaydi ✅",
            parse_mode="HTML"
        )
        return False


# === 📄 Yangi Konspekt ===
@router.message(F.text == "📄 Yangi Konspekt")
async def new_conspect(msg: types.Message):
    if is_blocked(msg.from_user.id):
        return await msg.answer("⛔ Siz bloklangansiz.")
    if not await check_limit(msg.from_user.id, msg): return
    await msg.answer("Fan nomini kiriting (masalan: Matematika):")
    set_state(msg.from_user.id, "subject")


# === 📘 Dars ishlanma yaratish ===
@router.message(F.text == "📘 Dars ishlanma yaratish")
async def new_lesson_plan_start(msg: types.Message):
    if is_blocked(msg.from_user.id):
        return await msg.answer("⛔ Siz bloklangansiz.")
    if not await check_limit(msg.from_user.id, msg): return
    await msg.answer("Fan nomini kiriting (masalan: Biologiya):")
    set_state(msg.from_user.id, "lesson_subject")


# === 📙 Metodik maslahat ===
@router.message(F.text == "📙 Metodik maslahat")
async def methodical_start(msg: types.Message):
    uid = msg.from_user.id
    if is_blocked(uid):
        return await msg.answer("⛔ Siz bloklangansiz.")
    if not await check_limit(uid, msg): return
    await msg.answer("Fan nomini kiriting (masalan: Tarix):")
    set_state(uid, "method_subject")


# === 🪄 Muammoni tahlil qilish ===
@router.message(F.text == "🪄 Muammoni tahlil qilish")
async def problem_analysis_start(msg: types.Message):
    uid = msg.from_user.id
    if is_blocked(uid):
        return await msg.answer("⛔ Siz bloklangansiz.")
    if not await check_limit(uid, msg): return
    await msg.answer(
        "🧩 Darsda duch kelgan muammoingizni yozing.\n\n"
        "Masalan:\n"
        "— O‘quvchilar mavzuni tushunmayapti.\n"
        "— Darsda vaqt yetmayapti.\n"
        "— Guruh ishlari sust kechadi va hokazo."
    )
    set_state(uid, "problem_text")


# === Asosiy text jarayoni ===
@router.message(F.text)
async def text_flow_handler(msg: types.Message):
    uid = msg.from_user.id
    state = get_state(uid)
    text = msg.text.strip()

    if state == "problem_text":
        if not await check_limit(uid, msg): return
        await msg.answer("⏳ Muammo tahlil qilinmoqda, biroz kuting...")
        try:
            result = analyze_teaching_problem(text)
            set_state(uid, None)
            return await msg.answer(result, reply_markup=main_menu())
        except Exception as e:
            set_state(uid, None)
            return await msg.answer(f"❌ Xatolik: {e}", reply_markup=main_menu())

    if state in ["subject", "lesson_subject", "method_subject"]:
        set_subject(uid, text)
        next_state = {"subject": "grade", "lesson_subject": "lesson_grade", "method_subject": "method_grade"}[state]
        set_state(uid, next_state)
        return await msg.answer("Sinfni kiriting (masalan: 8):")

    if state in ["grade", "lesson_grade", "method_grade"]:
        set_grade(uid, text)
        next_state = {"grade": "topic", "lesson_grade": "lesson_topic", "method_grade": "method_topic"}[state]
        set_state(uid, next_state)
        return await msg.answer("Endi mavzuni kiriting:")

    if state == "topic":
        if not await check_limit(uid, msg): return
        subject, grade, topic = get_subject(uid), get_grade(uid), text
        await msg.answer("⏳ Konspekt tayyorlanmoqda...")
        content = generate_conspect(subject, grade, topic)
        filename = create_named_docx(content, subject, topic, uid)
        save_history(uid, subject, grade, topic, filename)
        await msg.answer_document(types.FSInputFile(filename), caption="✅ Konspekt tayyor!", reply_markup=main_menu())
        os.remove(filename)
        set_state(uid, None)

    elif state == "lesson_topic":
        if not await check_limit(uid, msg): return
        subject, grade, topic = get_subject(uid), get_grade(uid), text
        await msg.answer("⏳ Dars ishlanma tayyorlanmoqda...")
        plan = generate_lesson_plan(subject, grade, topic)
        filename = create_named_docx(plan, subject, topic + "_DarsIshlanma", uid)
        save_history(uid, subject, grade, topic, filename)
        await msg.answer_document(types.FSInputFile(filename), caption="✅ Dars ishlanma tayyor!", reply_markup=main_menu())
        os.remove(filename)
        set_state(uid, None)

    elif state == "method_topic":
        if not await check_limit(uid, msg): return
        subject, grade, topic = get_subject(uid), get_grade(uid), text
        await msg.answer("⏳ Metodik maslahat tayyorlanmoqda...")
        result = generate_methodical_advice(subject, grade, topic)
        set_state(uid, None)
        return await msg.answer(result, reply_markup=main_menu())


# === 💳 To‘lov cheki yuborish ===
@router.message(F.photo)
async def handle_payment_photo(msg: types.Message):
    user_id = msg.from_user.id
    if is_blocked(user_id):
        return await msg.answer("⛔ Siz bloklangansiz.")
    username = msg.from_user.username or "Noma’lum"
    photo_id = msg.photo[-1].file_id
    payment_id = add_payment(user_id, username, photo_id)
    await msg.answer("✅ To‘lov cheki qabul qilindi! Admin tekshiradi ⏳")
    buttons = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="✅ Tasdiqlash", callback_data=f"approve_{payment_id}"),
            InlineKeyboardButton(text="❌ Rad etish", callback_data=f"reject_{payment_id}")
        ],
        [
            InlineKeyboardButton(text="📩 Foydalanuvchi bilan bog‘lanish",
                                 url=f"https://t.me/{username}" if username != "Noma’lum" else f"tg://user?id={user_id}")
        ]
    ])
    await msg.bot.send_photo(
        ADMIN_ID,
        photo=photo_id,
        caption=f"💳 <b>Yangi to‘lov!</b>\n\n👤 @{username}\n🆔 ID: <code>{user_id}</code>\n📎 Payment ID: <code>{payment_id}</code>",
        parse_mode="HTML",
        reply_markup=buttons
    )


# === 📤 Excel fayldan konspekt yaratish ===
@router.message(F.text.contains("Excel fayldan"))
async def excel_instruction(msg: types.Message):
    text = (
        "📘 Excel fayldan konspekt yaratish bo‘limi.\n\n"
        "🧩 Excel faylni quyidagicha tayyorlang:\n"
        "1. Faqat bitta ustun bo‘lsin — <b>Mavzu</b> (birinchi qatorda yozing).\n"
        "2. Quyidagicha ko‘rinishda bo‘lsin:\n\n"
        "   | Mavzu |\n"
        "   |----------------------------|\n"
        "   | Kasrlarni qo‘shish |\n"
        "   | Quyosh tizimi |\n"
        "   | Fe’l zamonlari |\n"
        "   | Kimyoviy reaksiyalar |\n\n"
        "3. Faylni .xlsx formatda saqlang.\n"
        "4. So‘ng faylni shu yerga yuboring 📎"
    )
    await msg.answer(text, parse_mode="HTML")
    await msg.answer("📎 Endi faylni yuboring (faqat .xlsx formatda).")
    set_state(msg.from_user.id, "excel_upload")


@router.message(F.document)
async def handle_excel_file(msg: types.Message):
    user_id = msg.from_user.id
    state = get_state(user_id)
    if state != "excel_upload":
        return
    document = msg.document
    file_path = f"temp_{user_id}.xlsx"
    await msg.bot.download(document, file_path)
    try:
        df = pd.read_excel(file_path)
        if df.empty or "Mavzu" not in df.columns:
            await msg.answer("❌ Fayl noto‘g‘ri. Excelda 'Mavzu' nomli ustun bo‘lishi kerak.")
            os.remove(file_path)
            return
        topics = df["Mavzu"].dropna().tolist()
        total = len(topics)
        await msg.answer(f"⏳ {total} ta mavzu uchun konspekt yaratilmoqda, biroz kuting...")
        topics_text = "\n".join([f"{i+1}. {t}" for i, t in enumerate(topics)])
        prompt = (
            f"Quyidagi {total} ta mavzu bo‘yicha o‘qituvchilar uchun juda uzun, batafsil konspekt yozing.\n"
            f"Har bir mavzuga alohida sarlavha qo‘ying, strukturani saqlang.\n\n{topics_text}"
        )
        result = generate_conspect("Umumiy fan", "Har xil sinflar", prompt)
        doc = Document()
        doc.add_heading("Yig‘ma Konspekt", level=0)
        doc.add_paragraph(result)
        output = f"{user_id}_yigma_konspekt.docx"
        doc.save(output)
        await msg.answer_document(types.FSInputFile(output), caption="✅ Yig‘ma konspekt tayyor!")
        os.remove(file_path)
        os.remove(output)
    except Exception as e:
        await msg.answer(f"❌ Xatolik: {str(e)}")
        if os.path.exists(file_path):
            os.remove(file_path)
